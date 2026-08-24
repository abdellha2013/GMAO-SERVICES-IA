"""
app/data_sources/file/docx_loader.py
====================================

Description
-----------
Loader spécialisé pour les fichiers Microsoft Word DOCX.

Ce loader :

- valide le fichier ;
- vérifie l'extension ;
- détecte le type MIME ;
- extrait le texte des paragraphes ;
- extrait le texte des tableaux ;
- construit un SourceDocument standardisé.

Le résultat est ensuite utilisé par le pipeline RAG :

    DOCX
      ↓
    DOCXLoader
      ↓
    SourceDocument
      ↓
    Parser
      ↓
    Chunker
      ↓
    Embedding
"""

from __future__ import annotations

from pathlib import Path
from typing import Final

from app.data_sources.file.file_source import FileSource
from app.models.document import SourceDocument

class DOCXLoader(FileSource):
    """
    Loader pour les fichiers Microsoft Word DOCX.

    Parameters
    ----------
    source:
        Chemin du fichier DOCX.
    """

    SUPPORTED_EXTENSIONS: Final[tuple[str, ...]] = (".docx",)

    def __init__(self, source: str | Path) -> None:
        """
        Initialise le loader DOCX.
        """

        super().__init__(source)

    # ==========================================================
    # Source Information
    # ==========================================================

    @property
    def source_name(self) -> str:
        """
        Retourne le nom du fichier source.
        """

        return self.filename

    # ==========================================================
    # Load
    # ==========================================================

    def load(self) -> SourceDocument:
        """
        Charge et extrait le contenu du fichier DOCX.

        Le contenu extrait comprend :

        - les paragraphes ;
        - les tableaux ;
        - les métadonnées du document.

        Returns
        -------
        SourceDocument
            Document standardisé.

        Raises
        ------
        FileValidationError
            Si le fichier n'est pas un DOCX valide.

        FileLoadingError
            Si le document ne peut pas être chargé.
        """

        from zipfile import BadZipFile

        from app.data_sources.file.validators import (
            ensure_non_empty_content,
            ensure_zip_based_format,
        )
        from app.exceptions import (
            FileLoadingError,
            InvalidDOCXError,
        )
                    


        self.logger.info(
            "Loading DOCX file '%s'.",
            self.path,
        )

        # ==========================================================
        # Validation
        # ==========================================================

        self.validate()

        self.ensure_extension(
            *self.SUPPORTED_EXTENSIONS
        )

        ensure_zip_based_format(self.path)

        # ==========================================================
        # Import dependency
        # ==========================================================

        try:
            from docx import Document as PythonDocument

        except ImportError as exc:
            raise FileLoadingError(
                message=(
                    "The 'python-docx' package is required "
                    "to load DOCX files."
                ),
                original=exc,
            ) from exc

        # ==========================================================
        # Load DOCX
        # ==========================================================

        try:
            document = PythonDocument(self.path)

        except BadZipFile as exc:

            self.logger.error(
                "DOCX file '%s' is not a valid Microsoft Word document.",
                self.path,
            )

            raise InvalidDOCXError(
                message=(
                    f"File '{self.filename}' is not a valid "
                    f"Microsoft Word document."
                ),
                original=exc,
            ) from exc

        except Exception as exc:

            self.logger.error(
                "Unable to open DOCX file '%s'.",
                self.path,
            )

            raise FileLoadingError(
                message=(
                    f"Unable to load DOCX file "
                    f"'{self.filename}'."
                ),
                original=exc,
            ) from exc

        # ==========================================================
        # Extract paragraphs
        # ==========================================================

        paragraphs: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        # ==========================================================
        # Extract tables
        # ==========================================================

        tables: list[str] = []

        for table in document.tables:
            rows: list[str] = []

            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                # Ignore completely empty rows
                if not any(cells):
                    continue

                rows.append(
                    " | ".join(cells)
                )

            if rows:
                tables.append(
                    "\n".join(rows)
                )

        # ==========================================================
        # Build content
        # ==========================================================

        sections: list[str] = []

        if paragraphs:
            sections.append(
                "\n".join(paragraphs)
            )

        if tables:
            sections.append(
                "\n\n".join(tables)
            )

        content = "\n\n".join(sections)

        # ==========================================================
        # Empty document
        # ==========================================================

        ensure_non_empty_content(
            content,
            self.filename,
            logger=self.logger,
        )

        # ==========================================================
        # Metadata
        # ==========================================================

        metadata = dict(
            self.metadata()
        )

        metadata.update(
            {
                "content_length": len(content),
                "paragraphs_count": len(paragraphs),
                "tables_count": len(document.tables),
                "table_rows_count": sum(
                    len(table.rows)
                    for table in document.tables
                ),
            }
        )

        # ==========================================================
        # Result
        # ==========================================================

        self.logger.info(
            "DOCX file '%s' loaded successfully.",
            self.filename,
        )

        return SourceDocument(
            source_name=self.filename,
            source_type="docx",
            source_path=self.path,
            content=content,
            mime_type=self.mime_type,
            size=self.size,
            created_at=self.created_at,
            updated_at=self.modified_at,
            metadata=metadata,
        )


    